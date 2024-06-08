from docx import Document
from docx.shared import Inches

def makePrintables(personDict, documentNumber=0):
    document = Document()
    r = document.add_paragraph().add_run()
    for key in personDict:
        r.add_picture(personDict[key].barcodePath, width=Inches(2.5))
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    document.save('barcodesheet' + str(documentNumber) + '.docx')