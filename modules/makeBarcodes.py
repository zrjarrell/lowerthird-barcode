import barcode
from barcode.writer import ImageWriter
from tkinter import filedialog
from docx import Document
from docx.shared import Inches

def makeBarcodes(personDict):
    barcodeDir = filedialog.askdirectory()
    upc = barcode.get_barcode_class('upc')
    for key in personDict:
        newBarcode = upc(key, writer=ImageWriter())
        newBarcode.render(text=getFullName(personDict[key])).save(barcodeDir + '/' + key + '.png')
        personDict[key]['barcodePath'] = barcodeDir + '/' + key + '.png'

def makePrintables(personDict, documentNumber=0):
    document = Document()
    r = document.add_paragraph().add_run()
    for key in personDict:
        r.add_picture(personDict[key]['barcodePath'], width=Inches(2.5))
    sections = document.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    document.save('barcodesheet' + str(documentNumber) + '.docx')

def getFullName(person):
    return person['firstName'] + ' ' + person['lastName']